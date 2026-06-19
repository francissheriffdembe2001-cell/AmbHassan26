from docx import Document
from docx.shared import Pt
from datetime import datetime


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def main():
    doc = Document()
    doc.core_properties.title = 'SLeClear MIS — Project Documentation'
    doc.core_properties.author = 'SLeClear Generator'
    doc.core_properties.created = datetime.now()

    add_heading(doc, 'SLeClear MIS — Project Overview', level=1)
    add_paragraph(doc, 'This document summarises the SLeClear MIS project, how to set it up locally, the database schema, and relevant files.')

    add_heading(doc, 'Quick Setup', level=2)
    add_paragraph(doc, 'Prerequisites: Python 3.8+, MySQL (XAMPP), and pip installed.')
    add_paragraph(doc, '1. Create virtual environment and install requirements:')
    add_paragraph(doc, '   python -m venv venv')
    add_paragraph(doc, '   venv\\Scripts\\activate (Windows)')
    add_paragraph(doc, '   python -m pip install -r requirements.txt')

    add_heading(doc, 'Database', level=2)
    add_paragraph(doc, 'The project ships with `database.sql` which creates `sleclear_db` with sample data. Two helper scripts are provided:')
    add_paragraph(doc, ' - `init_db.py`: runs SQL (useful for XAMPP default root user with no password).')
    add_paragraph(doc, ' - `create_db.py`: alternative loader using mysql-connector multi-statement execution.')
    add_paragraph(doc, 'phpMyAdmin: http://localhost/phpmyadmin — import `database.sql` or use the scripts.')

    add_heading(doc, 'App Configuration', level=2)
    add_paragraph(doc, 'Database connection is configured in `app.py`. You can set environment variables to override defaults:')
    add_paragraph(doc, ' - DB_HOST, DB_PORT, DB_USER, DB_PASS, DB_NAME, DB_CHARSET')
    add_paragraph(doc, 'Or provide a single `DATABASE_URL` in the form: mysql://user:pass@host:port/dbname')

    add_heading(doc, 'Default Credentials', level=2)
    add_paragraph(doc, 'The sample users and plain passwords (for development) are:')
    add_paragraph(doc, ' - admin / admin123')
    add_paragraph(doc, ' - finance / finance123')
    add_paragraph(doc, ' - registry / registry123')

    add_heading(doc, 'Important Files', level=2)
    add_paragraph(doc, ' - app.py : Flask application and DB access')
    add_paragraph(doc, ' - database.sql : schema + sample data')
    add_paragraph(doc, ' - init_db.py / create_db.py : database loaders')
    add_paragraph(doc, ' - templates/ and static/ : frontend templates and assets')

    add_heading(doc, 'How to Run', level=2)
    add_paragraph(doc, '1. Ensure MySQL/XAMPP is running.')
    add_paragraph(doc, '2. Initialise DB: python init_db.py')
    add_paragraph(doc, '3. Start the app: python app.py')
    add_paragraph(doc, '4. Open browser: http://localhost:5000')

    add_heading(doc, 'Security and Next Steps', level=2)
    add_paragraph(doc, ' - Change the MySQL root password and update env vars.')
    add_paragraph(doc, ' - Do not run Flask debug server in production. Use a WSGI server (gunicorn/uwsgi).')
    add_paragraph(doc, ' - Consider storing secrets using a .env file and a secrets manager.')

    add_heading(doc, 'Contact / Notes', level=2)
    add_paragraph(doc, 'Generated on: ' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    out_path = 'SLeClear_Project_Documentation.docx'
    doc.save(out_path)
    print('[OK] Generated', out_path)


if __name__ == '__main__':
    main()
